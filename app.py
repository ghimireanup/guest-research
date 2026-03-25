"""
Guest Research AI Agent - Flask Backend
- Gemini AI for brief, questions, category analysis, and timeline
- DuckDuckGo for web search, YouTube videos, and top articles
- Polling model (no SSE) for cloud compatibility
"""

import os
import json
import time
import threading
import uuid
import re
import io
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from bs4 import BeautifulSoup
from ddgs import DDGS
from google import genai
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_folder="static")
CORS(app)

GEMINI_MODEL = "gemini-2.5-flash"

_gemini_client = None

def get_gemini_client():
    global _gemini_client
    if _gemini_client is None:
        _gemini_client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
    return _gemini_client


# -------------------------------------------------------------------
# Session store
# -------------------------------------------------------------------

research_sessions: dict[str, dict] = {}


def new_session(guest_name: str, links: list, context: str) -> str:
    session_id = str(uuid.uuid4())
    research_sessions[session_id] = {
        "guest_name":  guest_name,
        "links":       links,
        "context":     context,
        "status":      "Starting research...",
        "step":        0,
        "sections":         {},   # all rendered sections live here
        "images":           [],
        "categories":       [],
        "industry_content": {},
        "brief":            None,
        "questions":        None,
        "done":             False,
        "error":            None,
        "created_at":       time.time(),
    }
    return session_id


def update_session(session_id: str, **kwargs):
    if session_id in research_sessions:
        research_sessions[session_id].update(kwargs)


def push_section(session_id: str, key: str, value):
    """Add or update a single key inside sessions[id]['sections']."""
    if session_id in research_sessions:
        research_sessions[session_id]["sections"][key] = value


# -------------------------------------------------------------------
# Image extraction
# -------------------------------------------------------------------

BLOCKED_IMAGE_PATTERNS = [
    "logo", "icon", "spinner", "pixel", "avatar", "badge",
    "button", "arrow", "1x1", "spacer", "blank", ".svg",
]

def extract_image_from_url(url: str) -> str | None:
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        resp = requests.get(url, headers=headers, timeout=8)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")

        for selector in [
            ("meta", {"property": "og:image"}),
            ("meta", {"name": "og:image"}),
            ("meta", {"name": "twitter:image"}),
        ]:
            tag = soup.find(*selector)
            if tag and tag.get("content"):
                img = tag["content"]
                if not any(p in img.lower() for p in BLOCKED_IMAGE_PATTERNS):
                    return img

        for img in soup.find_all("img", src=True):
            src = img["src"]
            if src.startswith("//"):  src = "https:" + src
            if not src.startswith("http"): continue
            if any(p in src.lower() for p in BLOCKED_IMAGE_PATTERNS): continue
            w = img.get("width", "")
            if str(w).isdigit() and int(w) < 100: continue
            return src
    except Exception as e:
        print(f"[image] {url}: {e}")
    return None


def collect_images(urls: list[str], max_images: int = 6) -> list[str]:
    images, seen = [], set()
    for url in urls:
        if len(images) >= max_images: break
        img = extract_image_from_url(url)
        if img and img not in seen:
            seen.add(img); images.append(img)
    return images


# -------------------------------------------------------------------
# Web search (DuckDuckGo text)
# -------------------------------------------------------------------

def search_web(query: str, max_results: int = 6) -> list[dict]:
    try:
        with DDGS() as ddgs:
            return list(ddgs.text(query, max_results=max_results))
    except Exception as e:
        print(f"[search] {e}")
        return []


# -------------------------------------------------------------------
# Page scraper
# -------------------------------------------------------------------

def scrape_page(url: str, max_chars: int = 3000) -> str:
    headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
    try:
        resp = requests.get(url, headers=headers, timeout=8)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "nav", "header", "footer", "aside", "form"]):
            tag.decompose()
        text = re.sub(r"\s+", " ", soup.get_text(separator=" ", strip=True))
        return text[:max_chars]
    except Exception as e:
        print(f"[scrape] {url}: {e}")
        return ""


# -------------------------------------------------------------------
# Step 1: Background research
# -------------------------------------------------------------------

def research_guest_background(guest_name: str, user_links: list[str]) -> tuple[str, list[str]]:
    scraped_urls, parts = [], []

    # User-provided links first (most reliable source)
    for url in user_links[:5]:
        page = scrape_page(url, 3000)
        if page:
            parts.append(f"[USER-PROVIDED: {url}]\n{page}")
            scraped_urls.append(url)

    # Web search supplement
    for query in [
        f"{guest_name} entrepreneur founder biography",
        f"{guest_name} company startup career achievements",
    ]:
        for result in search_web(query, 4)[:3]:
            if result.get("body"):
                parts.append(f"[SNIPPET] {result.get('title','')}: {result['body']}")
            if result.get("href") and result["href"] not in scraped_urls:
                page = scrape_page(result["href"], 2000)
                if page:
                    parts.append(f"[PAGE: {result['href']}]\n{page}")
                    scraped_urls.append(result["href"])

    seen, unique = set(), []
    for p in parts:
        if p not in seen: seen.add(p); unique.append(p)

    return "\n\n---\n\n".join(unique) or "No background found.", scraped_urls


# -------------------------------------------------------------------
# Step 2a: Past interviews
# -------------------------------------------------------------------

def find_guest_interviews(guest_name: str) -> tuple[str, list[str]]:
    parts, urls = [], []
    for query in [
        f"{guest_name} podcast interview",
        f"{guest_name} YouTube interview talk",
        f'"{guest_name}" podcast guest appearance',
    ]:
        for r in search_web(query, 5)[:4]:
            parts.append(
                f"Title: {r.get('title','')}\nURL: {r.get('href','')}\nSnippet: {r.get('body','')}"
            )
            if r.get("href"): urls.append(r["href"])
    return "\n\n---\n\n".join(parts) or "No interviews found.", urls


# -------------------------------------------------------------------
# Step 2b: Top YouTube videos (sorted by view count)
# -------------------------------------------------------------------

def find_youtube_videos(guest_name: str) -> list[dict]:
    """Search DuckDuckGo videos and return top 5 YouTube results by view count."""
    videos, seen = [], set()

    def _yt_thumb(url: str) -> str:
        m = re.search(r'(?:v=|youtu\.be/)([a-zA-Z0-9_-]{11})', url)
        return f"https://img.youtube.com/vi/{m.group(1)}/hqdefault.jpg" if m else ""

    # Primary: DuckDuckGo video search
    try:
        with DDGS() as ddgs:
            raw = list(ddgs.videos(f"{guest_name} interview talk podcast", max_results=20))
        for r in raw:
            url = r.get("content", "")
            if ("youtube.com/watch" not in url and "youtu.be/" not in url) or url in seen:
                continue
            seen.add(url)
            stats = r.get("statistics") or {}
            try: views = int(stats.get("viewCount") or 0)
            except: views = 0
            images = r.get("images") or {}
            videos.append({
                "title":       r.get("title", "Untitled"),
                "url":         url,
                "views":       views,
                "duration":    r.get("duration", ""),
                "thumbnail":   _yt_thumb(url) or images.get("large") or images.get("medium") or "",
                "channel":     r.get("uploader", ""),
                "published":   (r.get("published") or "")[:10],
                "description": (r.get("description") or "")[:160],
            })
    except Exception as e:
        print(f"[youtube-ddgs] {e}")

    # Fallback: text search for YouTube URLs
    if len(videos) < 3:
        for r in search_web(f"site:youtube.com {guest_name} interview", 8):
            url = r.get("href", "")
            if "youtube.com/watch" not in url or url in seen: continue
            seen.add(url)
            videos.append({
                "title":       r.get("title", "Untitled"),
                "url":         url,
                "views":       0,
                "duration":    "",
                "thumbnail":   _yt_thumb(url),
                "channel":     "",
                "published":   "",
                "description": (r.get("body") or "")[:160],
            })

    videos.sort(key=lambda x: x["views"], reverse=True)
    return videos[:5]


# -------------------------------------------------------------------
# Step 2c: Top 5 articles from search engines
# -------------------------------------------------------------------

SKIP_DOMAINS = {
    "youtube.com", "twitter.com", "x.com", "instagram.com",
    "facebook.com", "tiktok.com", "reddit.com",
}

def find_top_articles(guest_name: str) -> list[dict]:
    """Return top 5 quality articles (non-social-media) about the guest."""
    articles, seen_urls = [], set()

    for query in [
        f"{guest_name} entrepreneur profile story",
        f"{guest_name} interview article feature",
        f"{guest_name} startup company profile",
    ]:
        if len(articles) >= 5: break
        for r in search_web(query, 8):
            url = r.get("href", "")
            if not url or url in seen_urls: continue
            if any(d in url for d in SKIP_DOMAINS): continue
            seen_urls.add(url)
            domain = re.sub(r"^www\.", "", url.split("/")[2]) if len(url.split("/")) > 2 else url
            articles.append({
                "title":   r.get("title", "Untitled"),
                "url":     url,
                "snippet": (r.get("body") or "")[:200],
                "source":  domain,
            })
            if len(articles) >= 5: break

    return articles[:5]


# -------------------------------------------------------------------
# Step 3: Gemini — guest categorization (structured JSON for industry search)
# -------------------------------------------------------------------

def categorize_guest_with_gemini(guest_name: str, background: str) -> list[dict]:
    """Call Gemini to identify 4-6 structured categories the guest belongs to.
    Each category includes search queries for past/present/future content."""
    prompt = f"""Analyze {guest_name} based on this background data.

{background[:4000]}

Return ONLY a raw JSON array — no markdown fences, no preamble, no explanation.
Each element must have exactly these keys:
- id: snake_case string identifier for this category
- label: human-readable category name
- domain: broader industry or topic string
- past_query: a DuckDuckGo search query for historical content about this domain (e.g. "venture capital history evolution 2010s")
- present_query: a DuckDuckGo search query for the current state of this domain (e.g. "venture capital trends 2024 2025")
- future_query: a DuckDuckGo search query for where this domain is heading (e.g. "future of venture capital AI predictions")

Return minimum 4 categories, maximum 6. Categories must reflect {guest_name}'s actual industries, topics, and domains — not generic ones.

Example for Alex Hormozi:
[
  {{"id": "acquisition_entrepreneurship", "label": "Acquisition Entrepreneurship", "domain": "business acquisitions and portfolio companies", "past_query": "business acquisition strategies history small business buyouts", "present_query": "acquisition entrepreneurship trends 2024 portfolio model", "future_query": "future of acquisition entrepreneurship AI automation"}},
  ...
]"""

    try:
        response = get_gemini_client().models.generate_content(model=GEMINI_MODEL, contents=prompt).text
        # Strip any markdown fences
        cleaned = re.sub(r"```(?:json)?\s*", "", response).strip().rstrip("`")
        return json.loads(cleaned)
    except Exception as e:
        print(f"[categorize] Error: {e}")
        return []


def search_industry_content(categories: list[dict]) -> dict:
    """For each category, search DuckDuckGo for past/present/future articles + YouTube videos.
    Returns a nested dict keyed by category id."""
    result = {}

    for cat in categories:
        cat_id = cat.get("id", "unknown")
        label  = cat.get("label", cat_id)
        domain = cat.get("domain", label)
        articles, videos = [], []
        seen_urls = set()

        # Run past/present/future web searches
        for timeframe, query_key in [("past", "past_query"), ("present", "present_query"), ("future", "future_query")]:
            query = cat.get(query_key, "")
            if not query:
                continue
            try:
                for r in search_web(query, 5):
                    url = r.get("href", "")
                    if not url or url in seen_urls:
                        continue
                    seen_urls.add(url)
                    articles.append({
                        "title":     r.get("title", "Untitled"),
                        "url":       url,
                        "snippet":   (r.get("body") or "")[:180],
                        "timeframe": timeframe,
                    })
            except Exception as e:
                print(f"[industry-{cat_id}-{timeframe}] {e}")

        # Deduplicate and limit to 5 articles
        articles = articles[:5]

        # YouTube search for this category's domain
        try:
            yt_query = f"{domain} site:youtube.com"
            for r in search_web(yt_query, 8):
                url = r.get("href", "")
                if "youtube.com/watch" not in url or url in seen_urls:
                    continue
                seen_urls.add(url)
                videos.append({
                    "title":     r.get("title", "Untitled"),
                    "url":       url,
                    "snippet":   (r.get("body") or "")[:180],
                    "timeframe": "present",
                })
                if len(videos) >= 5:
                    break
        except Exception as e:
            print(f"[industry-{cat_id}-youtube] {e}")

        result[cat_id] = {
            "label":    label,
            "articles": articles,
            "videos":   videos,
        }

    return result


# -------------------------------------------------------------------
# Step 3 (legacy): Gemini — category + timeline analysis (returns structured JSON)
# -------------------------------------------------------------------

def analyze_guest(guest_name: str, background: str, interviews: str) -> dict:
    """Quick Gemini call: identify 4-6 categories + past/present/future timeline."""
    prompt = f"""Analyze {guest_name} based on this research data.

BACKGROUND:
{background[:3000]}

INTERVIEWS/MEDIA:
{interviews[:2000]}

Return ONLY a valid JSON object — no markdown, no explanation, no code fences:
{{
  "categories": ["Specific Category 1", "Specific Category 2", "Specific Category 3", "Specific Category 4"],
  "past": "2-3 sentences about who they were and what they were known for 5-10 years ago.",
  "present": "2-3 sentences about what they are currently doing and known for.",
  "future": "2-3 sentences about where they appear to be heading and what they represent for the future."
}}

For categories, be specific and insightful (e.g. 'Serial Entrepreneur', 'Angel Investor', 'Bestselling Author', 'Venture Capitalist', 'Product Visionary', 'Impact Investor'). Minimum 4, maximum 6 categories."""

    try:
        response = get_gemini_client().models.generate_content(model=GEMINI_MODEL, contents=prompt).text
        m = re.search(r'\{[\s\S]*\}', response)
        if m:
            return json.loads(m.group())
    except Exception as e:
        print(f"[analyze] {e}")

    return {
        "categories": ["Entrepreneur", "Business Leader", "Innovator", "Speaker"],
        "past": "Historical information is being compiled.",
        "present": "Currently active in their field.",
        "future": "Future trajectory is being analyzed.",
    }


# -------------------------------------------------------------------
# Step 3b: Gemini — full guest brief
# -------------------------------------------------------------------

def build_context_with_gemini(guest_name: str, background: str,
                               interviews: str, context: str) -> str:
    context_note = f'\nThe interviewer\'s research goal: "{context}"\nTailor this brief to serve that goal.\n' if context.strip() else ""

    prompt = f"""You are a professional podcast researcher preparing a guest brief.

Guest Name: {guest_name}
{context_note}
--- BACKGROUND DATA ---
{background[:8000]}

--- INTERVIEW / MEDIA DATA ---
{interviews[:6000]}

Write a comprehensive guest brief with EXACTLY these sections:

## Who They Are
2-3 paragraphs covering their identity and what makes them notable.

## What They've Built
Key companies, products, or projects and why they matter.

## What They Stand For
Core beliefs, philosophy, and mission.

## Past Interviews & Media Appearances
List each as: - **[Title / Platform]** — [One-sentence summary]

## What They've Already Talked About
Recurring themes across their past interviews.

## Unexplored Angles
3-5 angles no interviewer has explored yet — aligned with the interviewer's goal if provided."""

    return get_gemini_client().models.generate_content(model=GEMINI_MODEL, contents=prompt).text


# -------------------------------------------------------------------
# Step 4: Gemini — 15 questions
# -------------------------------------------------------------------

def generate_questions_with_gemini(guest_name: str, brief: str, context: str) -> str:
    context_note = f'\nThe interviewer\'s goal: "{context}"\nDesign the questions to serve this angle.\n' if context.strip() else ""

    prompt = f"""You are a world-class podcast host designing interview questions.

Guest: {guest_name}
{context_note}
Guest Brief:
{brief[:5000]}

Generate exactly 15 interview questions in three sections:

### Part 1: Background & Journey (5 Questions)
Personal story and pivotal moments.

### Part 2: Deep Insights (5 Questions)
Expertise, work philosophy, hard-won lessons.

### Part 3: Never-Been-Asked Questions (5 Questions)
Based on the Unexplored Angles — questions no interviewer has asked before.

Rules: specific to this guest, open-ended, numbered 1-5 per section.

End with an "Interview Notes" section: 2-3 tactical tips for the interviewer."""

    return get_gemini_client().models.generate_content(model=GEMINI_MODEL, contents=prompt).text


# -------------------------------------------------------------------
# docx builder
# -------------------------------------------------------------------

def build_docx_bytes(guest_name: str, brief: str, questions: str,
                     industry_content: dict | None = None) -> bytes:
    doc = Document()
    today = datetime.now().strftime("%B %d, %Y")
    t = doc.add_heading(f"Guest Research — {guest_name}", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(today)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    doc.add_heading("GUEST BRIEF", level=1)
    _add_markdown_to_doc(doc, brief)

    # Industry Context page (if available)
    if industry_content:
        doc.add_page_break()
        doc.add_heading("INDUSTRY CONTEXT", level=1)
        for cat_id, data in industry_content.items():
            doc.add_heading(data.get("label", cat_id), level=2)
            doc.add_heading("Articles", level=3)
            if data.get("articles"):
                for i, a in enumerate(data["articles"], 1):
                    tf = a.get("timeframe", "").capitalize()
                    _inline_bold(doc.add_paragraph(style="List Number"),
                                 f"({tf}) {a.get('title', '')} — {a.get('url', '')}")
            else:
                doc.add_paragraph("No articles found for this category.")
            doc.add_heading("YouTube Videos", level=3)
            if data.get("videos"):
                for i, v in enumerate(data["videos"], 1):
                    tf = v.get("timeframe", "").capitalize()
                    _inline_bold(doc.add_paragraph(style="List Number"),
                                 f"({tf}) {v.get('title', '')} — {v.get('url', '')}")
            else:
                doc.add_paragraph("No videos found for this category.")

    doc.add_page_break()
    doc.add_heading("INTERVIEW QUESTIONS", level=1)
    _add_markdown_to_doc(doc, questions)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_markdown_to_doc(doc, text: str):
    for line in text.split("\n"):
        line = line.strip()
        if not line:                     doc.add_paragraph()
        elif line.startswith("## "):     doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):    doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ","* ")):
            _inline_bold(doc.add_paragraph(style="List Bullet"), line[2:])
        elif re.match(r"^\d+\.\s", line):
            _inline_bold(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s*", "", line))
        else:
            _inline_bold(doc.add_paragraph(), line)


def _inline_bold(para, text: str):
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part: para.add_run(part).bold = (i % 2 == 1)


# -------------------------------------------------------------------
# Research pipeline — background thread
# -------------------------------------------------------------------

def run_research(session_id: str, guest_name: str, user_links: list[str], context: str):
    try:
        # ── Step 1: Background ───────────────────────────────────
        update_session(session_id, step=1, status=f"Researching {guest_name}...")
        background, scraped_urls = research_guest_background(guest_name, user_links)
        images = collect_images(scraped_urls[:8], max_images=6)
        update_session(session_id, images=images)

        # ── Step 2: All searches in parallel ────────────────────
        update_session(session_id, step=2, status="Finding interviews, YouTube videos, and articles...")

        with ThreadPoolExecutor(max_workers=3) as ex:
            f_interviews = ex.submit(find_guest_interviews, guest_name)
            f_youtube    = ex.submit(find_youtube_videos,   guest_name)
            f_articles   = ex.submit(find_top_articles,     guest_name)

        interviews, interview_urls = f_interviews.result()
        youtube_videos             = f_youtube.result()
        top_articles               = f_articles.result()

        # More images from interview pages
        if len(images) < 4:
            more = collect_images(interview_urls[:6], max_images=3)
            images = list(dict.fromkeys(images + more))[:6]
            update_session(session_id, images=images)

        # Push YouTube + articles immediately so UI can render them
        push_section(session_id, "youtube",   youtube_videos)
        push_section(session_id, "articles",  top_articles)

        # ── Step 3: Guest categorization + industry content ──────
        try:
            update_session(session_id, step=3, status="Analyzing industry context and finding top content...")
            categories = categorize_guest_with_gemini(guest_name, background)
            if categories:
                industry_content = search_industry_content(categories)
            else:
                industry_content = {}
            update_session(session_id, categories=categories, industry_content=industry_content)
        except Exception as e:
            print(f"[industry-step] Error (continuing): {e}")
            # Non-fatal — continue pipeline with empty data

        # ── Step 4: Gemini analysis ─────────────────────────────
        update_session(session_id, step=4, status="Identifying categories and timeline...")
        analysis = analyze_guest(guest_name, background, interviews)
        push_section(session_id, "categories", analysis.get("categories", []))
        push_section(session_id, "temporal",   {
            "past":    analysis.get("past", ""),
            "present": analysis.get("present", ""),
            "future":  analysis.get("future", ""),
        })

        update_session(session_id, status="Building comprehensive guest brief...")
        brief = build_context_with_gemini(guest_name, background, interviews, context)
        push_section(session_id, "brief", brief)

        # ── Step 5: Questions ────────────────────────────────────
        update_session(session_id, step=5, status="Generating 15 interview questions...")
        questions = generate_questions_with_gemini(guest_name, brief, context)
        push_section(session_id, "questions", questions)

        # ── Done ─────────────────────────────────────────────────
        update_session(session_id, step=6, status="Done!", done=True,
                       brief=brief, questions=questions)

    except Exception as e:
        print(f"[research] Error: {e}")
        import traceback; traceback.print_exc()
        update_session(session_id, done=True, error=str(e))


# -------------------------------------------------------------------
# Routes
# -------------------------------------------------------------------

@app.route("/")
def index():
    return send_from_directory("static", "index.html")


@app.route("/api/research", methods=["POST"])
def start_research():
    data       = request.get_json() or {}
    guest_name = data.get("guest_name", "").strip()
    user_links = [l.strip() for l in data.get("links", []) if l.strip()]
    context    = data.get("context", "").strip()

    if not guest_name:
        return jsonify({"error": "guest_name is required"}), 400
    if not os.getenv("GOOGLE_API_KEY"):
        return jsonify({"error": "GOOGLE_API_KEY is not set. Add it to your .env file."}), 500

    session_id = new_session(guest_name, user_links, context)
    threading.Thread(
        target=run_research,
        args=(session_id, guest_name, user_links, context),
        daemon=True,
    ).start()
    return jsonify({"session_id": session_id})


@app.route("/api/status/<session_id>")
def get_status(session_id: str):
    sess = research_sessions.get(session_id)
    if not sess:
        return jsonify({"error": "Session not found"}), 404

    return jsonify({
        "step":             sess["step"],
        "status":           sess["status"],
        "sections":         sess["sections"],
        "images":           sess["images"],
        "categories":       sess["categories"],
        "industry_content": sess["industry_content"],
        "done":             sess["done"],
        "error":            sess["error"],
        "download_ready":   sess["brief"] is not None and sess["questions"] is not None,
    })


@app.route("/api/download/<session_id>")
def download_file(session_id: str):
    sess = research_sessions.get(session_id)
    if not sess or not sess.get("brief") or not sess.get("questions"):
        return jsonify({"error": "Research not ready or session expired"}), 404

    docx_bytes = build_docx_bytes(sess["guest_name"], sess["brief"], sess["questions"],
                                  sess.get("industry_content", {}))
    safe = re.sub(r"[^\w\s-]", "", sess["guest_name"]).strip().replace(" ", "_")
    filename = f"Guest_Research_{safe}_{datetime.now().strftime('%Y-%m-%d')}.docx"

    return send_file(
        io.BytesIO(docx_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    print(f"\n🎙️  Guest Research Agent running at http://localhost:{port}\n")
    app.run(debug=True, port=port, threaded=True)
